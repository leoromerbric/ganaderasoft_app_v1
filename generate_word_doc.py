#!/usr/bin/env python3
"""
Script to generate a consolidated Word document from markdown documentation.
This script consolidates all documentation from the docs/ folder into a single Word document,
preserving the structure and including diagrams.
"""

import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import markdown2

# Constants
DOCS_DIR = Path(__file__).parent / "docs"
WORD_DIR = Path(__file__).parent / "Word"
OUTPUT_FILE = WORD_DIR / "GanaderaSoft_Documentacion_Consolidada.docx"

# Document order based on README.md
DOC_ORDER = [
    "README.md",
    "arquitectura.md",
    "estrategia-offline.md",
    "modulos.md",
    "base-datos.md",
    "api-servicios.md",
    "testing.md",
    "configuracion.md"
]


def add_hyperlink(paragraph, url, text):
    """
    A function that places a hyperlink within a paragraph object.
    """
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id, )

    # Create a new run object (a wrapper over a text element)
    new_run = OxmlElement('w:r')

    # Set the run's text
    rPr = OxmlElement('w:rPr')

    # Set the run's style to a hyperlink style
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    new_run.append(rPr)

    new_run.text = text

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    return hyperlink


def create_document():
    """Create and configure the Word document."""
    doc = Document()
    
    # Set up document properties
    doc.core_properties.title = "GanaderaSoft - Documentación Consolidada"
    doc.core_properties.author = "GanaderaSoft Team"
    doc.core_properties.comments = "Documentación técnica consolidada del sistema GanaderaSoft"
    
    return doc


def add_title_page(doc):
    """Add a title page to the document."""
    title = doc.add_heading("GanaderaSoft", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph("Documentación Técnica Consolidada")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(18)
    subtitle.runs[0].font.color.rgb = RGBColor(68, 68, 68)
    
    doc.add_paragraph()
    
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run("Sistema de gestión integral para fincas ganaderas\n").font.size = Pt(12)
    info.add_run("Desarrollado en Flutter").font.size = Pt(12)
    
    doc.add_page_break()


def add_table_of_contents(doc):
    """Add a table of contents placeholder."""
    heading = doc.add_heading("Índice de Contenidos", level=1)
    
    toc_paragraph = doc.add_paragraph()
    toc_paragraph.add_run("Esta sección contiene el índice de la documentación consolidada.").font.italic = True
    
    doc.add_paragraph()
    
    for i, doc_name in enumerate(DOC_ORDER, 1):
        if doc_name == "README.md":
            doc.add_paragraph(f"{i}. Visión General", style='List Number')
        else:
            title = doc_name.replace('.md', '').replace('-', ' ').title()
            doc.add_paragraph(f"{i}. {title}", style='List Number')
    
    doc.add_page_break()


def extract_mermaid_diagrams(content):
    """
    Extract mermaid diagrams from markdown content and replace them with placeholders.
    Returns the modified content and a list of diagrams.
    """
    diagrams = []
    pattern = r'```mermaid\n(.*?)\n```'
    
    def replace_diagram(match):
        diagram_code = match.group(1)
        diagrams.append(diagram_code)
        return f"\n[DIAGRAMA MERMAID #{len(diagrams)}]\n"
    
    modified_content = re.sub(pattern, replace_diagram, content, flags=re.DOTALL)
    return modified_content, diagrams


def add_mermaid_diagram(doc, diagram_code, diagram_number):
    """Add a mermaid diagram as a code block with clear labeling."""
    doc.add_paragraph(f"Diagrama #{diagram_number}:", style='Intense Quote')
    
    # Add the diagram code in a styled paragraph
    code_paragraph = doc.add_paragraph(diagram_code)
    code_paragraph.style = 'No Spacing'
    
    # Format as code
    for run in code_paragraph.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0, 0, 139)
    
    doc.add_paragraph()


def process_markdown_content(doc, content, file_name):
    """Process markdown content and add it to the document."""
    # Extract mermaid diagrams
    content, diagrams = extract_mermaid_diagrams(content)
    
    # Split content into lines for processing
    lines = content.split('\n')
    
    i = 0
    diagram_counter = 0
    
    while i < len(lines):
        line = lines[i]
        
        # Check for mermaid diagram placeholder
        if '[DIAGRAMA MERMAID #' in line:
            diagram_counter += 1
            if diagram_counter <= len(diagrams):
                add_mermaid_diagram(doc, diagrams[diagram_counter - 1], diagram_counter)
            i += 1
            continue
        
        # Process headers
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            text = line.lstrip('#').strip()
            if text:
                doc.add_heading(text, level=min(level, 9))
        
        # Process code blocks
        elif line.startswith('```'):
            code_lang = line[3:].strip()
            code_lines = []
            i += 1
            
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            
            if code_lines:
                # Add code block with language label
                if code_lang:
                    label = doc.add_paragraph(f"Código ({code_lang}):", style='Intense Quote')
                
                code_text = '\n'.join(code_lines)
                code_paragraph = doc.add_paragraph(code_text)
                
                # Format as code
                for run in code_paragraph.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0, 0, 139)
        
        # Process lists
        elif line.strip().startswith(('- ', '* ', '+ ')):
            text = line.strip()[2:].strip()
            if text:
                doc.add_paragraph(text, style='List Bullet')
        
        elif re.match(r'^\d+\.\s', line.strip()):
            text = re.sub(r'^\d+\.\s', '', line.strip())
            if text:
                doc.add_paragraph(text, style='List Number')
        
        # Process regular paragraphs
        elif line.strip():
            # Handle bold and italic
            paragraph = doc.add_paragraph()
            
            # Simple bold/italic processing
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = paragraph.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('*') and part.endswith('*'):
                    run = paragraph.add_run(part[1:-1])
                    run.italic = True
                elif part.startswith('`') and part.endswith('`'):
                    run = paragraph.add_run(part[1:-1])
                    run.font.name = 'Courier New'
                    run.font.color.rgb = RGBColor(220, 20, 60)
                else:
                    paragraph.add_run(part)
        
        i += 1


def process_documentation(doc):
    """Process all documentation files in order."""
    for doc_name in DOC_ORDER:
        doc_path = DOCS_DIR / doc_name
        
        if not doc_path.exists():
            print(f"Warning: {doc_name} not found, skipping...")
            continue
        
        print(f"Processing {doc_name}...")
        
        # Add section marker
        if doc_name != "README.md":
            doc.add_page_break()
        
        # Read the markdown file
        with open(doc_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Process the content
        process_markdown_content(doc, content, doc_name)
        
        print(f"Completed {doc_name}")


def main():
    """Main function to generate the Word document."""
    print("="*60)
    print("Generando documento Word consolidado de GanaderaSoft")
    print("="*60)
    print()
    
    # Create Word directory if it doesn't exist
    WORD_DIR.mkdir(exist_ok=True)
    print(f"Carpeta Word: {WORD_DIR}")
    
    # Create the document
    print("Creando documento Word...")
    doc = create_document()
    
    # Add title page
    print("Agregando página de título...")
    add_title_page(doc)
    
    # Add table of contents
    print("Agregando índice de contenidos...")
    add_table_of_contents(doc)
    
    # Process all documentation files
    print("Procesando archivos de documentación...")
    process_documentation(doc)
    
    # Save the document
    print(f"\nGuardando documento en: {OUTPUT_FILE}")
    doc.save(OUTPUT_FILE)
    
    print()
    print("="*60)
    print("✓ Documento generado exitosamente!")
    print(f"✓ Ubicación: {OUTPUT_FILE}")
    print("="*60)


if __name__ == "__main__":
    main()
